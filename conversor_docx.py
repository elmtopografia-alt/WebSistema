#!/usr/bin/env python3
"""
Conversor DOCX -> Estrutura JSON para SGT Propostas V2.1
CORREÇÃO: Preserva formatação de runs individuais (negrito, itálico, tamanho)
"""

import sys
import json
import re
import mammoth
from docx import Document
from docx.shared import Pt, Inches, RGBColor
import base64
import os
import importlib.util

# =====================================================
# CONFIGURAÇÃO DE CORES POR SEÇÃO (Ciclo Azul→Marrom→Cinza)
# =====================================================
CORES_SECAO = [
    {'cor': '#1e40af', 'borda': '#3b82f6', 'nome': 'Azul'},      # Seção 1
    {'cor': '#7c2d12', 'borda': '#92400e', 'nome': 'Marrom'},    # Seção 2
    {'cor': '#374151', 'borda': '#6b7280', 'nome': 'Cinza'},     # Seção 3
    {'cor': '#1e3a8a', 'borda': '#1d4ed8', 'nome': 'Azul Escuro'}, # Seção 4
    {'cor': '#92400e', 'borda': '#b45309', 'nome': 'Marrom Escuro'}, # Seção 5
    {'cor': '#4b5563', 'borda': '#9ca3af', 'nome': 'Cinza Claro'},  # Seção 6
    {'cor': '#2563eb', 'borda': '#60a5fa', 'nome': 'Azul Claro'},   # Seção 7
    {'cor': '#78350f', 'borda': '#a16207', 'nome': 'Marrom Claro'}, # Seção 8
    {'cor': '#1f2937', 'borda': '#4b5563', 'nome': 'Cinza Escuro'}, # Seção 9
    {'cor': '#1d4ed8', 'borda': '#3b82f6', 'nome': 'Azul Médio'},   # Seção 10
]

# Variáveis do sistema que devem ser sempre incluídas
VARIAVEIS_SISTEMA = [
    'cidade_limpo',      # Cidade sem estado (para cabeçalho)
    'DataExtenso',       # Data por extenso
    'numero_proposta',   # Número da proposta
    'whatsapp',          # WhatsApp da empresa
]

# =====================================================
# FUNÇÕES DE EXTRAÇÃO DE ESTILOS (CORRIGIDAS V2.1)
# =====================================================

def extrair_estilos_run(run):
    """
    Extrai estilos CSS de um run específico (trecho de texto formatado).
    CORREÇÃO V2.1: Agora verifica run.font diretamente para negrito, itálico, tamanho.
    """
    estilos = {}

    if run.font:
        font = run.font

        # Tamanho da fonte (prioridade máxima: run individual)
        if font.size and font.size.pt:
            estilos['font-size'] = f"{font.size.pt}pt"

        # NEGRITO - verifica propriedade direta do run (CORREÇÃO PRINCIPAL)
        if font.bold:
            estilos['font-weight'] = 'bold'

        # Itálico
        if font.italic:
            estilos['font-style'] = 'italic'

        # Sublinhado
        if font.underline:
            estilos['text-decoration'] = 'underline'

        # Cor da fonte
        if font.color and font.color.rgb:
            estilos['color'] = f"#{font.color.rgb}"

        # Fonte family (se disponível)
        if font.name:
            estilos['font-family'] = font.name

    return estilos


def extrair_estilos_paragrafo(paragraph):
    """
    Extrai estilos do parágrafo (alinhamento, espaçamento, etc.).
    NÃO inclui formatação de caracteres (negrito, itálico) - isso vem dos runs.
    """
    estilos = {}

    # Alinhamento do parágrafo
    alignment_map = {
        0: 'left', 1: 'center', 2: 'right', 3: 'justify'
    }
    if paragraph.alignment is not None:
        estilos['text-align'] = alignment_map.get(paragraph.alignment, 'left')

    # Tamanho base do parágrafo (fallback se run não tiver tamanho)
    if paragraph.style and paragraph.style.font:
        font = paragraph.style.font
        if font.size and font.size.pt:
            estilos['font-size-base'] = f"{font.size.pt}pt"
        if font.bold:
            estilos['font-weight-base'] = 'bold'
        if font.italic:
            estilos['font-style-base'] = 'italic'

    return estilos


def processar_paragrafo_com_runs(paragraph):
    """
    Processa parágrafo preservando formatação de cada run individual.
    Retorna lista de partes com seus estilos específicos.
    """
    partes = []
    estilos_para = extrair_estilos_paragrafo(paragraph)

    for run in paragraph.runs:
        texto = run.text
        if not texto:
            continue

        # Extrai estilos específicos deste run (CORREÇÃO V2.1)
        estilos_run = extrair_estilos_run(run)

        # Merge: run tem prioridade sobre parágrafo
        estilos_finais = {**estilos_para, **estilos_run}

        # Se run não tem tamanho definido, usa o tamanho base do parágrafo
        if 'font-size' not in estilos_finais and 'font-size-base' in estilos_finais:
            estilos_finais['font-size'] = estilos_finais.pop('font-size-base')

        partes.append({
            'texto': texto,
            'estilos': estilos_finais,
            'negrito': estilos_finais.get('font-weight') == 'bold',
            'italico': estilos_finais.get('font-style') == 'italic'
        })

    return partes


def gerar_html_runs(partes):
    """
    Gera HTML inline preservando formatação de cada run individualmente.
    Ex: texto normal <strong>${Empresa}</strong> texto normal
    """
    html = ''
    for parte in partes:
        texto = parte['texto']
        negrito = parte.get('negrito', False)
        italico = parte.get('italico', False)
        sublinhado = parte['estilos'].get('text-decoration') == 'underline'

        # Aplica tags inline apenas onde o run realmente tem a formatação
        if sublinhado:
            texto = f'<u>{texto}</u>'
        if italico:
            texto = f'<em>{texto}</em>'
        if negrito:
            texto = f'<strong>{texto}</strong>'

        html += texto
    return html


def extrair_estilos_unificados(paragraph):
    """
    Extrai estilos do parágrafo e gera HTML inline preservando formatação por run.
    CORREÇÃO: Só aplica font-weight:bold no bloco se TODOS os runs forem negrito.
    Caso contrário, usa conteudo_html com <strong> apenas onde necessário.
    """
    partes = processar_paragrafo_com_runs(paragraph)

    if not partes:
        return extrair_estilos_paragrafo(paragraph), "", ""

    # Texto plano (sem formatação)
    texto_completo = ''.join([p['texto'] for p in partes])

    # HTML com formatação inline por run
    html_completo = gerar_html_runs(partes)

    # Obtém estilos base do parágrafo
    estilos_para = extrair_estilos_paragrafo(paragraph)

    # Negrito apenas se TODOS os runs forem negrito (parágrafo realmente todo negrito)
    todos_negrito = all(p['negrito'] for p in partes) or estilos_para.get('font-weight-base') == 'bold'
    algum_negrito = any(p['negrito'] for p in partes)
    tem_italico = all(p['italico'] for p in partes) or estilos_para.get('font-style-base') == 'italic'

    # Tamanho mais frequente entre os runs
    tamanho_principal = None
    tamanhos_runs = [p['estilos'].get('font-size') for p in partes if p['estilos'].get('font-size')]
    if tamanhos_runs:
        from collections import Counter
        tamanho_principal = Counter(tamanhos_runs).most_common(1)[0][0]

    # Monta estilos CSS do bloco (parágrafo como um todo)
    estilos_finais = extrair_estilos_paragrafo(paragraph)

    # Aplica negrito ao bloco apenas se for universal no parágrafo
    if todos_negrito:
        estilos_finais['font-weight'] = 'bold'
    if tem_italico:
        estilos_finais['font-style'] = 'italic'
    if tamanho_principal:
        estilos_finais['font-size'] = tamanho_principal

    # Flag para indicar que o HTML já tem formatação inline mista
    estilos_finais['_tem_formatacao_mista'] = algum_negrito and not todos_negrito

    # Limpa chaves temporárias
    estilos_finais.pop('font-size-base', None)
    estilos_finais.pop('font-weight-base', None)
    estilos_finais.pop('font-style-base', None)

    return estilos_finais, texto_completo, html_completo


# =====================================================
# FUNÇÕES AUXILIARES (mantidas da V2.0)
# =====================================================

def _celula_e_negrita(cell):
    """Verifica se a célula tem texto e todas as partes com conteúdo são negrito."""
    for para in cell.paragraphs:
        if not para.text.strip():
            continue
        for run in para.runs:
            if run.text.strip() and not run.bold:
                return False
    return True


def _linha_e_cabecalho(row, row_index):
    """
    Detecta se uma linha da tabela é cabeçalho.
    Critério 1: Estilo Word indica 'tblHeader' (linha de cabeçalho real).
    Critério 2: Todas as células com conteúdo são negrito (heurística).
    """
    # Critério 1: Verifica atributo XML de cabeçalho de tabela
    try:
        ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        trPr = row._tr.find(f'{{{ns}}}trPr')
        if trPr is not None:
            tblHeader = trPr.find(f'{{{ns}}}tblHeader')
            if tblHeader is not None:
                return True
    except Exception:
        pass

    # Critério 2: Só aplica heurística de "todos negrito" na PRIMEIRA linha
    # (para não tratar erroneamente linhas de dados que sejam bold)
    if row_index == 0:
        celulas_com_texto = [c for c in row.cells if c.text.strip()]
        if celulas_com_texto and all(_celula_e_negrita(c) for c in celulas_com_texto):
            return True

    return False


def processar_tabela(table):
    """Converte tabela DOCX em estrutura HTML/CSS com detecção real de cabeçalho."""
    rows = []
    for row_index, row in enumerate(table.rows):
        cells = []
        seen_ids = set()

        # Detecta se esta linha é realmente um cabeçalho
        is_header = _linha_e_cabecalho(row, row_index)

        for cell in row.cells:
            # Evita células duplicadas por colspan
            cell_id = id(cell._tc)
            if cell_id in seen_ids:
                continue
            seen_ids.add(cell_id)

            # Lê colspan do XML
            try:
                ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
                tcPr = cell._tc.find(f'{{{ns}}}tcPr')
                grid_span = 1
                if tcPr is not None:
                    gs = tcPr.find(f'{{{ns}}}gridSpan')
                    if gs is not None:
                        grid_span = int(gs.get(f'{{{ns}}}val', 1))
            except Exception:
                grid_span = 1

            # Extrai texto e detecta negrito da célula
            textos = []
            celula_negrita = True  # Presume negrito até achar run sem negrito
            tem_texto = False

            for para in cell.paragraphs:
                if para.text.strip():
                    textos.append(para.text)
                    tem_texto = True
                    for run in para.runs:
                        if run.text.strip() and not run.bold:
                            celula_negrita = False

            if not tem_texto:
                celula_negrita = False

            cells.append({
                'texto': ' '.join(textos),
                'colspan': grid_span,
                'negrito': celula_negrita,
                'is_header_cell': is_header or celula_negrita,
                'estilos': {
                    'border': '1px solid #dee2e6',
                    'padding': '12px 15px',
                    'background': '#f8f9fa' if (is_header or celula_negrita) else 'transparent',
                    'vertical-align': 'top',
                    'font-weight': 'bold' if celula_negrita else 'normal'
                }
            })

        rows.append({
            'celulas': cells,
            'is_header': is_header,
            # Mantém 'linhas' no formato legado para compatibilidade
            '__celulas_legado': cells
        })

    # Converte para formato de saída compatível com o sistema legado
    # Cada item de rows agora é um dict com 'celulas' e 'is_header'
    linhas_saida = []
    for row_data in rows:
        linhas_saida.append(row_data['celulas'])

    return {
        'tipo': 'tabela',
        'linhas': linhas_saida,
        'linhas_meta': [{'is_header': r['is_header']} for r in rows],
        'estilos': {
            'width': '100%',
            'border-collapse': 'collapse',
            'margin': '25px 0',
            'font-size': '14px'
        }
    }


def detectar_variaveis(texto):
    """Detecta ${var}, ${ var }, {{var}} ou {{ var }} no texto"""
    padrao = r'\$\{\s*(\w+)\s*\}|\{\{\s*(\w+)\s*\}\}'
    matches = re.findall(padrao, texto)
    return list(set([m[0] or m[1] for m in matches if m[0] or m[1]]))


def detectar_variaveis_implicitas(texto):
    """Detecta variáveis que devem ser adicionadas implicitamente baseado no contexto"""
    vars_implicitas = []
    texto_lower = texto.lower()

    # Se menciona cidade + data, provavelmente precisa de cidade_limpo
    if any(x in texto_lower for x in ['cidade', 'data', 'extenso']) and '${cidade' in texto:
        vars_implicitas.append('cidade_limpo')

    # Se tem número de proposta
    if 'proposta' in texto_lower and ('nº' in texto_lower or 'numero' in texto_lower):
        vars_implicitas.append('numero_proposta')

    return vars_implicitas


def processar_imagens(doc, docx_path):
    """Extrai imagens embutidas e converte para base64"""
    imagens = []
    rels = doc.part.rels

    for rel in rels.values():
        if "image" in rel.target_ref:
            try:
                image = rel.target_part.blob
                ext = rel.target_ref.split('.')[-1]
                base64_img = base64.b64encode(image).decode('utf-8')
                imagens.append({
                    'id': rel.rId,
                    'extensao': ext,
                    'base64': f"data:image/{ext};base64,{base64_img}",
                    'tamanho': len(image)
                })
            except Exception as e:
                pass

    return imagens


def classificar_bloco(texto, estilos, nivel_titulo):
    """Classifica o bloco por conteúdo e contexto"""
    texto_lower = texto.strip().lower()

    # Se já detectou como título pelo nível, prioriza
    if nivel_titulo > 0:
        return 'titulo'

    # Detecção por conteúdo específico
    if any(x in texto_lower for x in ['cliente', 'contratante', 'nome:', 'e-mail:', 'email:']):
        return 'dados_cliente'
    elif any(x in texto_lower for x in ['obra', 'local', 'terreno', 'endereço:', 'bairro:', 'cidade/estado:']):
        return 'dados_obra'
    elif any(x in texto_lower for x in ['valor', 'investimento', 'r$', 'preço', 'mobilização', 'restante']):
        return 'valores'
    elif any(x in texto_lower for x in ['prazos', 'dias úteis', 'cronograma', 'etapa', 'mobilização']):
        return 'prazos'
    elif any(x in texto_lower for x in ['pagamento', 'banco:', 'pix:', 'agência:', 'conta:']):
        return 'pagamento'
    elif any(x in texto_lower for x in ['equipamento', 'vante', 'drone', 'gps', 'rtk', 'estação total', 'veículo']):
        return 'equipamentos'
    elif estilos.get('font-weight') == 'bold' and len(texto) < 120:
        return 'titulo'
    else:
        return 'texto_geral'


def detectar_nivel_titulo(para):
    """
    Detecta nível de título H1/H2/H3.
    Hierarquia:
      H1 → Título do documento (Title), Heading 1, seções principais (8. Prazos)
      H2 → Heading 2, subseções (8.1 Detalhe), FASE X
      H3 → Heading 3, sub-subseções (8.1.1 ...)
    """
    texto = para.text.strip()
    if not texto:
        return 0

    style_name = para.style.name if para.style else ''

    # 1. Estilos oficiais do Word (PT + EN)
    if style_name in ('Title', 'Título', 'title', 'título'):
        return 1
    # Subtítulo do documento → H2
    if style_name in ('Subtitle', 'Subtítulo', 'subtitle', 'subtítulo'):
        return 2
    if 'Heading 1' in style_name or 'Título 1' in style_name or 'Cabeçalho 1' in style_name:
        return 1
    if 'Heading 2' in style_name or 'Título 2' in style_name or 'Cabeçalho 2' in style_name:
        return 2
    if 'Heading 3' in style_name or 'Título 3' in style_name or 'Cabeçalho 3' in style_name:
        return 3

    # 2. Heurísticas por numeração manual
    # Sub-subseções "8.1.1 Item" → H3 (verificar antes de X.X)
    if re.match(r'^\d+\.\d+\.\d+\s+', texto):
        return 3
    # Subseções "8.1 Detalhe" → H2
    if re.match(r'^\d+\.\d+\s+', texto):
        return 2
    # Seções principais "8. Prazos Estimados" → H1 (título de seção)
    if re.match(r'^\d+\.\s+[A-ZÀ-Ú]', texto):
        return 1
    # FASE X → H2
    if re.match(r'^FASE\s+\d+', texto, re.IGNORECASE):
        return 2

    # 3. Fallback: todo em maiúsculas e curto = H1
    if texto.isupper() and len(texto) < 100:
        return 1

    return 0


def gerar_css_dinamico(blocos):
    """Gera CSS dinâmico com cores alternadas por seção de título H2"""
    css_base = """
        .modelo-docx { font-family: 'Segoe UI', Arial, sans-serif; line-height: 1.6; color: #333; max-width: 900px; margin: 0 auto; }

        /* Título principal H1 */
        .modelo-docx h1 { 
            color: #1e3a8a; 
            border-bottom: 3px solid #1e3a8a; 
            padding-bottom: 10px; 
            font-size: 24px; 
            font-weight: bold; 
            text-align: center;
        }

        /* Subtítulos H3 */
        .modelo-docx h3 { 
            color: #4b5563; 
            font-size: 14px; 
            font-weight: bold; 
            margin-top: 20px; 
        }

        /* Classes utilitárias */
        .titulo-principal { color: #1e3a8a; border-bottom: 3px solid #1e3a8a; padding-bottom: 10px; }
        .titulo-secao { color: #1e40af; margin-top: 25px; }
        .var-placeholder { background: #dbeafe; color: #1e40af; padding: 2px 6px; border-radius: 4px; font-family: monospace; border: 1px dashed #3b82f6; }
        .tabela-proposta th { background: #f1f5f9; font-weight: 600; }
        .dados_cliente, .dados_obra { background: #f8fafc; padding: 15px; border-radius: 8px; margin: 10px 0; }

        /* Cabeçalho e Rodapé vindos do DOCX */
        .sgt-texto-header_footer { 
            font-size: 11px !important; 
            color: #64748b !important; 
            text-align: center !important;
            margin: 0 !important;
            padding: 2px 0 !important;
        }
        .docx-header { border-bottom: 1px solid #e2e8f0; margin-bottom: 30px; padding-bottom: 15px; }
        .docx-footer { border-top: 1px solid #e2e8f0; margin-top: 50px; padding-top: 15px; }

        /* CORREÇÃO V2.1: Preservar negrito e tamanho de fonte inline */
        .sgt-texto-destaque { font-weight: bold !important; }
    """

    # Conta apenas os H2 para aplicar cores alternadas
    contador_h2 = 0
    css_secoes = []

    for bloco in blocos:
        if bloco.get('nivel_titulo') == 2:
            cor = CORES_SECAO[contador_h2 % len(CORES_SECAO)]
            css_secoes.append(f"""
        /* Seção {contador_h2 + 1} - {cor['nome']} */
        .modelo-docx h2:nth-of-type({contador_h2 + 1}) {{ 
            color: {cor['cor']}; 
            margin-top: 25px; 
            font-size: 18px; 
            font-weight: bold; 
            border-left: 4px solid {cor['borda']}; 
            padding-left: 12px; 
        }}""")
            contador_h2 += 1

    return css_base + '\n'.join(css_secoes)


def processar_secao_especial(secao_docx, variaveis_globais):
    """Processa parágrafos de uma seção especial (Header/Footer)"""
    blocos = []
    for para in secao_docx.paragraphs:
        texto = para.text
        if not texto.strip():
            continue

        vars_para = detectar_variaveis(texto)
        vars_implicitas = detectar_variaveis_implicitas(texto)
        todas_vars = list(set(vars_para + vars_implicitas))
        variaveis_globais.update(todas_vars)

        estilos, texto_extraido, html_runs = extrair_estilos_unificados(para)
        tem_mista = estilos.pop('_tem_formatacao_mista', False)

        blocos.append({
            'tipo': 'texto',
            'subtipo': 'header_footer',
            'conteudo': texto,
            'conteudo_html': html_runs if tem_mista else None,
            'estilos_css': estilos,
            'variaveis': todas_vars,
            'nivel_titulo': 0
        })

    for table in secao_docx.tables:
        tabela_proc = processar_tabela(table)
        texto_tabela = ' '.join([' '.join([c['texto'] for c in row]) for row in tabela_proc['linhas']])
        vars_tabela = detectar_variaveis(texto_tabela)
        variaveis_globais.update(vars_tabela)
        tabela_proc['variaveis'] = vars_tabela
        tabela_proc['subtipo'] = 'header_footer_table'
        blocos.append(tabela_proc)

    return blocos


# =====================================================
# FUNÇÃO PRINCIPAL DE CONVERSÃO (CORRIGIDA V2.1)
# =====================================================

def converter_docx(caminho_docx):
    """
    Função principal de conversão V2.1
    CORREÇÃO: Preserva formatação de runs individuais (negrito, itálico, tamanho)
    """

    if not os.path.exists(caminho_docx):
        return {'sucesso': False, 'erro': f'Arquivo nao encontrado: {caminho_docx}'}

    doc = Document(caminho_docx)

    # Extrai imagens primeiro
    imagens = processar_imagens(doc, caminho_docx)

    blocos = []
    variaveis_globais = set(VARIAVEIS_SISTEMA)  # Inicia com variáveis do sistema

    # Extração de Header e Footer (NOVO)
    blocos_header = []
    blocos_footer = []

    for section in doc.sections:
        if section.header:
            blocos_header.extend(processar_secao_especial(section.header, variaveis_globais))
        if section.footer:
            blocos_footer.extend(processar_secao_especial(section.footer, variaveis_globais))

    # =====================================================
    # CORREÇÃO V2.2: Processamento de parágrafos com runs
    # + Detecção de H1 consecutivo (título + subtítulo)
    # =====================================================
    ultimo_nivel = 0          # Nível do último bloco detectado
    blocos_sem_texto = 0      # Parágrafos de texto entre dois títulos

    for element in doc.element.body:
        if element.tag.endswith('p'):
            para = next((p for p in doc.paragraphs if p._element == element), None)
            if para:
                nivel = detectar_nivel_titulo(para)

                # ── Heurística Título + Subtítulo ────────────────────────────
                # Se o parágrafo anterior era H1 e este também seria H1,
                # e não houve texto entre eles → é o subtítulo do doc (H2)
                if nivel == 1 and ultimo_nivel == 1 and blocos_sem_texto == 0:
                    nivel = 2

                # Atualiza rastreadores
                if nivel > 0:
                    ultimo_nivel = nivel
                    blocos_sem_texto = 0
                # ─────────────────────────────────────────────────────────────

                # Extrai estilos + texto plano + HTML com runs preservados
                estilos, texto, html_runs = extrair_estilos_unificados(para)

                # Flag de formatação mista (ex: só ${Empresa} é negrito)
                tem_mista = estilos.pop('_tem_formatacao_mista', False)

                if not texto.strip() and nivel == 0:
                    blocos_sem_texto += 1
                    continue

                vars_para = detectar_variaveis(texto)
                vars_implicitas = detectar_variaveis_implicitas(texto)
                todas_vars = list(set(vars_para + vars_implicitas))
                variaveis_globais.update(todas_vars)

                tipo_bloco = classificar_bloco(texto, estilos, nivel)

                blocos.append({
                    'tipo': 'texto',
                    'subtipo': tipo_bloco,
                    'conteudo': texto,
                    # conteudo_html: HTML com <strong>/<em> apenas nos runs que têm formatação
                    # Usado pelo editor quando a formatação é mista (não todo o parágrafo)
                    'conteudo_html': html_runs if tem_mista else None,
                    'estilos_css': estilos,
                    'variaveis': todas_vars,
                    'nivel_titulo': nivel
                })

        elif element.tag.endswith('tbl'):
            table = next((t for t in doc.tables if t._element == element), None)
            if table:
                tabela_proc = processar_tabela(table)
                texto_tabela = ' '.join([' '.join([c['texto'] for c in row]) for row in tabela_proc['linhas']])
                vars_tabela = detectar_variaveis(texto_tabela)
                vars_imp_tabela = detectar_variaveis_implicitas(texto_tabela)
                todas_vars_tabela = list(set(vars_tabela + vars_imp_tabela))
                variaveis_globais.update(todas_vars_tabela)
                tabela_proc['variaveis'] = todas_vars_tabela

                blocos.append(tabela_proc)

    # Une os blocos (Header -> Body -> Footer)
    todos_blocos = blocos_header + blocos + blocos_footer

    # Gera CSS dinâmico baseado nos blocos processados
    css_dinamico = gerar_css_dinamico(todos_blocos)

    # Gera HTML preservando formatação
    try:
        with open(caminho_docx, "rb") as docx_file:
            result_mammoth = mammoth.convert_to_html(docx_file)
            html_base = result_mammoth.value

            # Se mammoth funcionou, adicionamos o header e footer manualmente no preview
            if blocos_header or blocos_footer:
                html_header = gerar_html_blocos(blocos_header, [])
                html_footer = gerar_html_blocos(blocos_footer, [])
                html_base = f'<header class="docx-header">{html_header}</header>' + html_base + f'<footer class="docx-footer">{html_footer}</footer>'
    except:
        html_base = gerar_html_blocos(todos_blocos, imagens)

    return {
        'sucesso': True,
        'nome_arquivo': os.path.basename(caminho_docx),
        'blocos': todos_blocos,
        'total_blocos': len(todos_blocos),
        'variaveis': sorted(list(variaveis_globais)),
        'total_variaveis': len(variaveis_globais),
        'imagens': len(imagens),
        'html_preview': html_base,
        'css_geral': css_dinamico
    }


def gerar_html_blocos(blocos, imagens):
    """Gera HTML preservando formatação original fallback"""
    html_parts = []

    for bloco in blocos:
        if bloco['tipo'] == 'texto':
            tag = 'p'
            classes = [bloco['subtipo']]

            if bloco['nivel_titulo'] == 1:
                tag = 'h1'
                classes.append('titulo-principal')
            elif bloco['nivel_titulo'] == 2:
                tag = 'h2'
                classes.append('titulo-secao')
            elif bloco['nivel_titulo'] == 3:
                tag = 'h3'

            if bloco['estilos_css'].get('font-weight') == 'bold':
                classes.append('sgt-texto-destaque')

            # Filtra chaves internas que não são CSS
            estilos_css = {k: v for k, v in bloco['estilos_css'].items() if not k.startswith('_')}
            estilos_inline = '; '.join([f"{k}:{v}" for k, v in estilos_css.items()])

            # Usa conteudo_html (com <strong>/<em> por run) se disponível
            conteudo = bloco.get('conteudo_html') or bloco['conteudo']

            # Substitui variáveis por placeholders visuais
            for var in bloco['variaveis']:
                conteudo = conteudo.replace(f"${{{var}}}", f'<span class="var-placeholder" data-var="{var}">{{{{{var}}}}}</span>')
                conteudo = conteudo.replace(f"{{{var}}}", f'<span class="var-placeholder" data-var="{var}">{{{{{var}}}}}</span>')

            html_parts.append(f'<{tag} class="{" ".join(classes)}" style="{estilos_inline}">{conteudo}</{tag}>')

        elif bloco['tipo'] == 'tabela':
            html_tbl = ['<table class="tabela-proposta" border="1" style="width:100%;border-collapse:collapse;margin:20px 0;">']
            for i, row in enumerate(bloco['linhas']):
                tag_row = 'th' if i == 0 else 'td'
                html_tbl.append('<tr>')
                for cell in row:
                    estilos = ';'.join([f"{k}:{v}" for k,v in cell['estilos'].items()])
                    texto_celula = cell["texto"].replace('\n', '<br>')
                    html_tbl.append(f'<{tag_row} colspan="{cell["colspan"]}" style="{estilos}">{texto_celula}</{tag_row}>')
                html_tbl.append('</tr>')
            html_tbl.append('</table>')
            html_parts.append(''.join(html_tbl))

    return '\n'.join(html_parts)


# =====================================================
# ENTRY POINT
# =====================================================

if __name__ == "__main__":
    # Configura o terminal para UTF-8 (Vital para Windows)
    if sys.platform == "win32":
        import io
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

    if len(sys.argv) > 1 and sys.argv[1] == '--test':
        print(json.dumps({
            "sucesso": True,
            "mensagem": "Python e dependências OK (V2.1 - Correção Negrito)",
            "python_version": sys.version,
            "mammoth": "OK" if importlib.util.find_spec("mammoth") else "Faltando",
            "docx": "OK" if importlib.util.find_spec("docx") else "Faltando"
        }, ensure_ascii=False))
        sys.exit(0)

    if len(sys.argv) < 2:
        print(json.dumps({
            "sucesso": False,
            "erro": "Uso: python conversor_docx.py <arquivo.docx> ou --test"
        }, ensure_ascii=False))
        sys.exit(1)

    try:
        resultado = converter_docx(sys.argv[1])
        print(json.dumps(resultado, ensure_ascii=False, indent=2))
    except Exception as e:
        import traceback
        print(json.dumps({
            'erro': str(e), 
            'traceback': traceback.format_exc(),
            'sucesso': False
        }))
